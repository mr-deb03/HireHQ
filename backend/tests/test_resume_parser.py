"""Unit tests for resume text extraction and the deterministic parser."""

from __future__ import annotations

import pytest

from app.core.exceptions import FileTooLarge, UnsupportedFileType
from app.modules.resumes.extraction import extract_text_sync
from app.providers.ai.heuristic import HeuristicAIProvider, tfidf_cosine
from app.providers.scanning import ScanVerdict, inspect_structure, validate_upload
from app.utils.skills import display_skill, extract_skills, match_skill, normalise_skill
from app.utils.text import extract_emails, extract_github, extract_linkedin, extract_phones
from tests.conftest import SAMPLE_RESUME, build_docx


@pytest.fixture
def provider() -> HeuristicAIProvider:
    return HeuristicAIProvider()


class TestTextExtraction:
    def test_extracts_docx(self):
        result = extract_text_sync(build_docx(SAMPLE_RESUME), "docx")
        assert result.is_usable
        assert "Rahul Sharma" in result.text
        assert "Flipkart" in result.text
        assert result.word_count > 50

    def test_docx_tables_are_read_row_wise(self):
        """Resume templates often put history in a table; column order must survive."""
        from io import BytesIO

        import docx

        document = docx.Document()
        document.add_paragraph("EXPERIENCE")
        table = document.add_table(rows=1, cols=3)
        cells = table.rows[0].cells
        cells[0].text = "Senior Engineer"
        cells[1].text = "Acme Corp"
        cells[2].text = "Jan 2020 - Present"
        buffer = BytesIO()
        document.save(buffer)

        result = extract_text_sync(buffer.getvalue(), "docx")
        line = next(ln for ln in result.text.splitlines() if "Acme Corp" in ln)
        assert "Senior Engineer" in line and "Jan 2020" in line

    def test_rejects_a_non_docx_masquerading_as_one(self):
        with pytest.raises(UnsupportedFileType):
            extract_text_sync(b"this is not a docx", "docx")

    def test_unknown_extension_is_rejected(self):
        with pytest.raises(UnsupportedFileType):
            extract_text_sync(b"data", "xyz")

    def test_near_empty_document_is_flagged_not_silently_accepted(self):
        result = extract_text_sync(build_docx("Hi"), "docx")
        assert not result.is_usable
        assert result.warnings


class TestUploadValidation:
    def test_accepts_a_valid_docx(self):
        extension, content_type = validate_upload(
            filename="cv.docx",
            content=build_docx(SAMPLE_RESUME),
            allowed_extensions={"pdf", "docx"},
            max_size_mb=10,
        )
        assert extension == "docx"
        assert "wordprocessingml" in content_type

    def test_rejects_disallowed_extension(self):
        with pytest.raises(UnsupportedFileType, match="not accepted"):
            validate_upload(
                filename="virus.exe",
                content=b"MZ\x90\x00",
                allowed_extensions={"pdf", "docx"},
                max_size_mb=10,
            )

    def test_rejects_content_that_contradicts_its_extension(self):
        """A .pdf whose bytes are not a PDF must be refused before anything reads it."""
        with pytest.raises(UnsupportedFileType, match="does not look like"):
            validate_upload(
                filename="resume.pdf",
                content=b"<html><script>alert(1)</script></html>",
                allowed_extensions={"pdf", "docx"},
                max_size_mb=10,
            )

    def test_rejects_oversized_file(self):
        with pytest.raises(FileTooLarge):
            validate_upload(
                filename="big.pdf",
                content=b"%PDF-" + b"0" * (3 * 1024 * 1024),
                allowed_extensions={"pdf"},
                max_size_mb=2,
            )

    def test_rejects_empty_file(self):
        with pytest.raises(UnsupportedFileType, match="empty"):
            validate_upload(
                filename="empty.pdf", content=b"", allowed_extensions={"pdf"}, max_size_mb=10
            )


class TestStructuralScanning:
    def test_clean_pdf_has_no_findings(self):
        assert inspect_structure(b"%PDF-1.4\nsimple content\n", "pdf") == []

    def test_detects_javascript_in_a_pdf(self):
        findings = inspect_structure(b"%PDF-1.4\n/JavaScript (app.alert(1))", "pdf")
        assert any("JavaScript" in f for f in findings)

    def test_detects_launch_action(self):
        findings = inspect_structure(b"%PDF-1.4\n/Launch /F (cmd.exe)", "pdf")
        assert any("launch action" in f for f in findings)

    def test_detects_eicar_signature(self):
        from app.providers.scanning import EICAR

        assert any("EICAR" in f for f in inspect_structure(b"%PDF-1.4" + EICAR, "pdf"))

    def test_unscanned_file_is_not_reported_as_clean(self):
        """Without a real AV engine the verdict must be NOT_SCANNED, never CLEAN."""
        import asyncio

        from app.providers.scanning import MalwareScanner

        result = asyncio.run(
            MalwareScanner().scan(build_docx("hello world " * 20), extension="docx")
        )
        assert result.verdict == ScanVerdict.NOT_SCANNED
        assert result.is_safe_to_store
        assert "not been scanned" in (result.detail or "")


class TestFieldExtraction:
    def test_email(self):
        assert extract_emails("Reach me at rahul.sharma@example.test today") == [
            "rahul.sharma@example.test"
        ]

    def test_phone(self):
        assert extract_phones("Phone: +91 9876543210")

    def test_year_is_not_mistaken_for_a_phone_number(self):
        assert extract_phones("Graduated in 2020") == []

    def test_linkedin(self):
        assert (
            extract_linkedin("profile: linkedin.com/in/rahulsharma")
            == "https://linkedin.com/in/rahulsharma"
        )

    def test_github(self):
        assert (
            extract_github("code at github.com/rahulsharma")
            == "https://github.com/rahulsharma"
        )

    def test_bare_github_domain_is_not_a_profile(self):
        assert extract_github("see github.com for details") is None


class TestSkillNormalisation:
    @pytest.mark.parametrize(
        ("raw", "canonical"),
        [
            ("ReactJS", "react"),
            ("React.js", "react"),
            ("REACT", "react"),
            ("TS", "typescript"),
            ("Postgres", "postgresql"),
            ("AWS", "amazon web services"),
            ("k8s", "kubernetes"),
            ("RESTful APIs", "rest api"),
        ],
    )
    def test_aliases_collapse(self, raw, canonical):
        assert normalise_skill(raw) == canonical

    @pytest.mark.parametrize(
        ("raw", "shown"),
        [
            ("javascript", "JavaScript"),
            ("typescript", "TypeScript"),
            ("rest api", "REST API"),
            ("reactjs", "React"),
            ("postgres", "PostgreSQL"),
            ("ci/cd", "CI/CD"),
            ("c++", "C++"),
        ],
    )
    def test_display_preserves_real_casing(self, raw, shown):
        assert display_skill(raw) == shown

    def test_extracts_skills_from_prose(self):
        found = extract_skills(
            "Built applications with React, TypeScript and PostgreSQL on AWS."
        )
        canonical = {normalise_skill(s) for s in found}
        assert {"react", "typescript", "postgresql", "amazon web services"} <= canonical

    def test_does_not_match_inside_unrelated_words(self):
        """'Go' must not match 'going', or every resume gains a Go skill."""
        found = {normalise_skill(s) for s in extract_skills("We are going forward")}
        assert "go" not in found


class TestSkillMatching:
    def test_exact(self):
        assert match_skill("React", ["React"]).matched

    def test_alias(self):
        match = match_skill("React", ["ReactJS"])
        assert match.matched and match.evidence == "ReactJS"

    def test_unrelated_does_not_match(self):
        assert not match_skill("React", ["PHP", "WordPress"]).matched

    def test_empty_profile_does_not_match(self):
        assert not match_skill("React", []).matched

    def test_near_miss_is_reported_as_missing(self):
        """Being generous here would inflate scores and mislead recruiters."""
        match = match_skill("Kubernetes", ["Kubectl"])
        assert not match.matched


class TestResumeParsing:
    async def test_parses_a_complete_resume(self, provider):
        result = await provider.parse_resume(text=SAMPLE_RESUME)
        parsed = result.value

        assert parsed.name == "Rahul Sharma"
        assert parsed.email == "rahul.sharma@example.test"
        assert parsed.phone
        assert parsed.linkedin_url and "rahulsharma" in parsed.linkedin_url
        assert parsed.github_url
        assert len(parsed.experience) == 2
        assert parsed.education
        assert parsed.confidence > 0.6

    async def test_extracts_current_role(self, provider):
        parsed = (await provider.parse_resume(text=SAMPLE_RESUME)).value
        current = next(e for e in parsed.experience if e.is_current)
        assert current.company == "Flipkart"
        assert "Frontend Engineer" in current.position

    async def test_computes_experience_from_dates(self, provider):
        parsed = (await provider.parse_resume(text=SAMPLE_RESUME)).value
        # Mar 2020 -> today; at minimum the two listed roles are contiguous.
        assert parsed.total_experience_years >= 5

    async def test_overlapping_roles_are_not_double_counted(self, provider):
        text = """Jane Doe
jane@example.test

EXPERIENCE
Engineer at A    Jan 2020 - Jan 2024
- Did work
Consultant at B    Jan 2020 - Jan 2024
- Did other work
"""
        parsed = (await provider.parse_resume(text=text)).value
        assert parsed.total_experience_years == pytest.approx(4.0, abs=0.2)

    async def test_reports_missing_fields_rather_than_inventing_them(self, provider):
        parsed = (await provider.parse_resume(text="Just a name\n\nSKILLS\nPython")).value
        assert "email" in parsed.missing_fields
        assert parsed.confidence < 0.6

    async def test_flags_a_scanned_document(self, provider):
        parsed = (await provider.parse_resume(text="   ")).value
        assert parsed.warnings

    async def test_never_infers_protected_attributes(self, provider):
        """The schema has no field for them, and none may appear in the output."""
        text = SAMPLE_RESUME + "\nDate of Birth: 12/03/1995\nGender: Male\nMarried"
        parsed = (await provider.parse_resume(text=text)).value
        dumped = parsed.model_dump()
        for forbidden in ("gender", "age", "date_of_birth", "marital_status",
                          "religion", "nationality"):
            assert forbidden not in dumped

    async def test_role_titles_containing_to_are_parsed(self, provider):
        """'Automation' and 'Practo' contain 'to'; the range splitter must not break."""
        text = """Deepa Chandran
deepa@example.test

EXPERIENCE
QA Automation Engineer at Practo    Jan 2021 - Present
- Built end-to-end test suites
"""
        parsed = (await provider.parse_resume(text=text)).value
        assert len(parsed.experience) == 1
        assert parsed.experience[0].company == "Practo"
        assert parsed.total_experience_years >= 4

    async def test_prose_under_experience_is_not_a_phantom_job(self, provider):
        text = """Fresh Graduate
grad@example.test

EXPERIENCE
No professional experience yet - recent graduate.

EDUCATION
B.Tech from IIT
"""
        parsed = (await provider.parse_resume(text=text)).value
        assert parsed.experience == []
        assert parsed.total_experience_years == 0


class TestJobDescriptionAnalysis:
    async def test_extracts_requirements(self, provider):
        result = await provider.analyze_job_description(
            title="Senior React Developer",
            description=(
                "We need a senior engineer.\n\n"
                "Requirements:\n"
                "- 5+ years of experience\n"
                "- Strong React and TypeScript\n"
                "- REST API design\n"
                "- Bachelor's degree in Computer Science\n\n"
                "Nice to have:\n"
                "- AWS and Docker\n"
            ),
        )
        analysis = result.value
        required = {normalise_skill(s.name) for s in analysis.required_skills}
        preferred = {normalise_skill(s.name) for s in analysis.preferred_skills}

        assert "react" in required
        assert "typescript" in required
        assert analysis.min_experience_years == 5
        assert "Bachelor's degree" in analysis.education_requirements
        assert "amazon web services" in preferred or "docker" in preferred

    async def test_detects_seniority(self, provider):
        result = await provider.analyze_job_description(
            title="Junior Developer", description="An entry level role for graduates."
        )
        assert result.value.seniority in ("JUNIOR", "INTERN")

    async def test_thin_description_reports_low_confidence(self, provider):
        result = await provider.analyze_job_description(
            title="Developer", description="We need a developer."
        )
        assert result.value.confidence < 0.6

    async def test_engine_is_labelled_honestly(self, provider):
        result = await provider.analyze_job_description(title="Dev", description="x" * 50)
        assert result.usage.engine == "heuristic-v1"
        assert provider.is_real_model is False


class TestSemanticSimilarity:
    def test_identical_text_is_maximally_similar(self):
        text = "Senior React developer with TypeScript and REST API experience"
        assert tfidf_cosine(text, text) > 0.99

    def test_unrelated_text_is_dissimilar(self):
        assert (
            tfidf_cosine(
                "Senior React developer with TypeScript",
                "Experienced pastry chef specialising in French desserts",
            )
            < 0.1
        )

    def test_related_text_scores_between(self):
        score = tfidf_cosine(
            "Senior React developer building web applications with TypeScript",
            "Frontend engineer experienced in React and TypeScript web development",
        )
        assert 0.15 < score < 0.99

    def test_empty_input_is_zero(self):
        assert tfidf_cosine("", "anything") == 0.0

    def test_is_symmetric(self):
        a, b = "React TypeScript developer", "TypeScript React engineer"
        assert tfidf_cosine(a, b) == pytest.approx(tfidf_cosine(b, a))

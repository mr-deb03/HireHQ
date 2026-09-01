"""Text extraction from resume documents.

PDF via ``pypdf`` and DOCX via ``python-docx``. Both run in a worker thread because they
are CPU-bound and would otherwise block the event loop for the duration of a large file.

Extraction is deliberately layout-aware where it matters: a DOCX's tables are flattened
row by row (many resume templates put the entire work history in a table, and reading it
cell-by-cell in document order scrambles it), and PDF pages are joined with explicit
separators so a section heading at the top of page 2 is not glued onto the last line of
page 1.
"""

from __future__ import annotations

import asyncio
import re
from dataclasses import dataclass
from io import BytesIO

from app.core.exceptions import UnsupportedFileType
from app.core.logging import get_logger

logger = get_logger(__name__)

#: Guards against a malformed or hostile document consuming the worker.
MAX_PAGES = 40
MAX_CHARS = 500_000

_LIGATURES = {
    "ﬀ": "ff", "ﬁ": "fi", "ﬂ": "fl", "ﬃ": "ffi", "ﬄ": "ffl",
    "‘": "'", "’": "'", "“": '"', "”": '"',
    "–": "-", "—": "-", " ": " ", "​": "",
    "•": "- ", "●": "- ", "▪": "- ", "": "- ",
}


@dataclass(slots=True)
class ExtractionResult:
    text: str
    page_count: int | None
    word_count: int
    #: Non-fatal problems worth surfacing, e.g. "this looks like a scanned image".
    warnings: list[str]

    @property
    def is_usable(self) -> bool:
        return self.word_count >= 30


def _clean(text: str) -> str:
    for source, replacement in _LIGATURES.items():
        text = text.replace(source, replacement)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Many PDFs emit one space between every glyph pair; collapse runs but keep single
    # spaces and paragraph breaks intact.
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()[:MAX_CHARS]


def _extract_pdf(content: bytes) -> ExtractionResult:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    warnings: list[str] = []
    try:
        reader = PdfReader(BytesIO(content), strict=False)
    except PdfReadError as exc:
        raise UnsupportedFileType(f"The PDF could not be read: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        # Many resumes are "encrypted" with an empty owner password, which pypdf can open.
        try:
            if reader.decrypt("") == 0:
                raise UnsupportedFileType(
                    "This PDF is password-protected. Please upload an unprotected copy."
                )
        except UnsupportedFileType:
            raise
        except Exception as exc:
            raise UnsupportedFileType(
                "This PDF is password-protected. Please upload an unprotected copy."
            ) from exc

    pages = reader.pages
    page_count = len(pages)
    if page_count > MAX_PAGES:
        warnings.append(f"Only the first {MAX_PAGES} of {page_count} pages were read.")
        pages = pages[:MAX_PAGES]

    chunks: list[str] = []
    for index, page in enumerate(pages):
        try:
            chunks.append(page.extract_text() or "")
        except Exception as exc:
            logger.warning("pdf_page_extract_failed", page=index, error=str(exc))
            warnings.append(f"Page {index + 1} could not be read.")

    text = _clean("\n\n".join(c for c in chunks if c.strip()))
    words = len(text.split())
    if words < 30:
        warnings.append(
            "Almost no text could be extracted. This is usually a scanned or "
            "image-based PDF, which needs OCR that this server does not perform."
        )
    return ExtractionResult(text=text, page_count=page_count, word_count=words, warnings=warnings)


def _extract_docx(content: bytes) -> ExtractionResult:
    import zipfile

    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(BytesIO(content))
    except (PackageNotFoundError, zipfile.BadZipFile, KeyError, ValueError) as exc:
        # A .docx is a zip; a truncated or mislabelled upload surfaces as BadZipFile
        # (or a KeyError for a missing part). Turn all of those into a clear 415 rather
        # than letting them escape as a 500.
        raise UnsupportedFileType(
            "The file is not a valid .docx document. If it is an older .doc file, "
            "please save it as .docx or PDF."
        ) from exc

    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

    # Resume templates frequently put the whole history in a table; read it row-wise so
    # "Company | Role | 2019-2022" stays on one line instead of being split apart.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Merged cells repeat their text across the row; collapse the repetition.
            deduped: list[str] = []
            for cell in cells:
                if cell and (not deduped or deduped[-1] != cell):
                    deduped.append(cell)
            if deduped:
                parts.append("  ".join(deduped))

    # Headers and footers often carry contact details.
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text.strip())

    text = _clean("\n".join(parts))
    words = len(text.split())
    warnings: list[str] = []
    if words < 30:
        warnings.append("Very little text could be extracted from this document.")
    return ExtractionResult(text=text, page_count=None, word_count=words, warnings=warnings)


def extract_text_sync(content: bytes, extension: str) -> ExtractionResult:
    extension = extension.lower().lstrip(".")
    if extension == "pdf":
        return _extract_pdf(content)
    if extension == "docx":
        return _extract_docx(content)
    if extension in ("txt", "text"):
        text = _clean(content.decode("utf-8", errors="replace"))
        return ExtractionResult(text, None, len(text.split()), [])
    raise UnsupportedFileType(f"Cannot extract text from a .{extension} file")


async def extract_text(content: bytes, extension: str) -> ExtractionResult:
    """Extract text off the event loop."""
    return await asyncio.to_thread(extract_text_sync, content, extension)

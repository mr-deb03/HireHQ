"""Development seed data.

Creates a realistic, self-consistent demo tenant: a super admin, one company with its
hiring team, jobs with real requirements, candidates with real resumes, applications
that have been scored by the actual ATS engine, interviews, feedback, offers and
workflows.

Nothing here is faked at the persistence layer - applications go through the real intake
service, resumes through the real parser, and scores through the real engine. That means
the seed doubles as an end-to-end smoke test of the whole pipeline.

Run with:  python -m app.db.seed  [--reset]
"""

from __future__ import annotations

import argparse
import asyncio
import random
import sys
import uuid
from datetime import UTC, date, datetime, timedelta

from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.config import settings
from app.core.enums import (
    ApplicationSource,
    ApplicationStatus,
    CompanySize,
    CompanyStatus,
    EmailTemplateKey,
    EmploymentType,
    InterviewRecommendation,
    InterviewType,
    JobStatus,
    RoleName,
    ScreeningQuestionType,
    SkillImportance,
    UserStatus,
    WorkflowActionType,
    WorkflowTrigger,
    WorkMode,
)
from app.core.logging import configure_logging, get_logger
from app.core.security import hash_password
from app.db.bootstrap import bootstrap_database, ensure_schema
from app.db.session import session_scope
from app.models.company import Company, CompanyLocation, Department
from app.models.job import Job, JobHiringTeamMember, JobScreeningQuestion, JobSkill
from app.models.user import Role, User
from app.models.workflow import Workflow, WorkflowStep
from app.utils.text import slugify

logger = get_logger("seed")

# Deterministic, so re-running the seed produces the same demo every time.
RANDOM = random.Random(20240501)

COMPANY_NAME = "Northwind Technologies"

# --------------------------------------------------------------------- people
STAFF: list[dict] = [
    {
        "email": "priya.nair@northwind.test",
        "first_name": "Priya",
        "last_name": "Nair",
        "job_title": "Head of Talent",
        "roles": [RoleName.COMPANY_ADMIN, RoleName.RECRUITER],
    },
    {
        "email": "arjun.mehta@northwind.test",
        "first_name": "Arjun",
        "last_name": "Mehta",
        "job_title": "Senior Recruiter",
        "roles": [RoleName.RECRUITER],
    },
    {
        "email": "sneha.rao@northwind.test",
        "first_name": "Sneha",
        "last_name": "Rao",
        "job_title": "Technical Recruiter",
        "roles": [RoleName.RECRUITER],
    },
    {
        "email": "vikram.desai@northwind.test",
        "first_name": "Vikram",
        "last_name": "Desai",
        "job_title": "Engineering Manager",
        "roles": [RoleName.HIRING_MANAGER],
    },
    {
        "email": "meera.krishnan@northwind.test",
        "first_name": "Meera",
        "last_name": "Krishnan",
        "job_title": "Director of Product",
        "roles": [RoleName.HIRING_MANAGER],
    },
    {
        "email": "rohan.gupta@northwind.test",
        "first_name": "Rohan",
        "last_name": "Gupta",
        "job_title": "Staff Engineer",
        "roles": [RoleName.INTERVIEWER, RoleName.EMPLOYEE],
    },
    {
        "email": "anita.shah@northwind.test",
        "first_name": "Anita",
        "last_name": "Shah",
        "job_title": "Senior Frontend Engineer",
        "roles": [RoleName.INTERVIEWER, RoleName.EMPLOYEE],
    },
    {
        "email": "karthik.iyer@northwind.test",
        "first_name": "Karthik",
        "last_name": "Iyer",
        "job_title": "Backend Architect",
        "roles": [RoleName.INTERVIEWER, RoleName.EMPLOYEE],
    },
    {
        "email": "divya.menon@northwind.test",
        "first_name": "Divya",
        "last_name": "Menon",
        "job_title": "Data Engineering Lead",
        "roles": [RoleName.INTERVIEWER, RoleName.EMPLOYEE],
    },
    {
        "email": "sameer.khan@northwind.test",
        "first_name": "Sameer",
        "last_name": "Khan",
        "job_title": "Product Designer",
        "roles": [RoleName.INTERVIEWER, RoleName.EMPLOYEE],
    },
]

DEPARTMENTS = ["Engineering", "Product", "Design", "Data", "Operations"]

LOCATIONS = [
    {"name": "Bengaluru HQ", "city": "Bengaluru", "state": "Karnataka", "country": "India",
     "timezone": "Asia/Kolkata", "is_headquarters": True},
    {"name": "Pune Office", "city": "Pune", "state": "Maharashtra", "country": "India",
     "timezone": "Asia/Kolkata", "is_headquarters": False},
    {"name": "Remote - India", "city": "Remote", "country": "India",
     "timezone": "Asia/Kolkata", "is_headquarters": False},
]

# ----------------------------------------------------------------------- jobs
JOBS: list[dict] = [
    {
        "title": "Senior React Developer",
        "department": "Engineering",
        "location": "Bengaluru, Karnataka",
        "work_mode": WorkMode.HYBRID,
        "employment_type": EmploymentType.FULL_TIME,
        "min_exp": 4, "max_exp": 8,
        "salary": (2200000, 3400000),
        "openings": 3,
        "status": JobStatus.PUBLISHED,
        "description": (
            "We are looking for a Senior React Developer to lead frontend work on our "
            "customer platform, which serves several million users.\n\n"
            "Responsibilities:\n"
            "- Build and maintain reusable React components and a shared design system\n"
            "- Develop and consume REST APIs alongside the backend team\n"
            "- Own frontend performance, accessibility and Core Web Vitals\n"
            "- Mentor junior developers and review their pull requests\n"
            "- Collaborate with product and design on scoping and delivery\n\n"
            "Requirements:\n"
            "- 4+ years of professional experience building web applications\n"
            "- Deep expertise in React, TypeScript and modern JavaScript\n"
            "- Strong experience with REST API integration and state management\n"
            "- Proficiency with Git and CI/CD workflows\n"
            "- Bachelor's degree in Computer Science or equivalent experience\n\n"
            "Nice to have:\n"
            "- Experience with Next.js and server-side rendering\n"
            "- Familiarity with AWS and Docker\n"
            "- Exposure to GraphQL\n"
            "- Testing with Jest and Playwright"
        ),
        "required": ["React", "TypeScript", "REST API", "Git", "JavaScript", "CSS"],
        "preferred": ["Next.js", "AWS", "Docker", "GraphQL"],
        "education": ["Bachelor's degree"],
        "responsibilities": [
            "Build and maintain reusable React components and a shared design system",
            "Develop and consume REST APIs alongside the backend team",
            "Own frontend performance, accessibility and Core Web Vitals",
            "Mentor junior developers and review their pull requests",
            "Collaborate with product and design on scoping and delivery",
        ],
    },
    {
        "title": "Backend Engineer (Python)",
        "department": "Engineering",
        "location": "Bengaluru, Karnataka",
        "work_mode": WorkMode.HYBRID,
        "employment_type": EmploymentType.FULL_TIME,
        "min_exp": 3, "max_exp": 7,
        "salary": (2000000, 3200000),
        "openings": 2,
        "status": JobStatus.PUBLISHED,
        "description": (
            "Join our platform team building the APIs and services behind Northwind.\n\n"
            "Responsibilities:\n"
            "- Design and build REST APIs using Python and FastAPI\n"
            "- Model and optimise PostgreSQL schemas and queries\n"
            "- Implement background processing with Redis and task queues\n"
            "- Write automated tests and maintain CI pipelines\n"
            "- Participate in on-call and incident response\n\n"
            "Requirements:\n"
            "- 3+ years of backend development experience\n"
            "- Strong Python skills and experience with FastAPI or Django\n"
            "- Solid PostgreSQL and SQL knowledge\n"
            "- Experience with Docker and REST API design\n"
            "- Bachelor's degree in a technical field\n\n"
            "Preferred:\n"
            "- Kubernetes and AWS experience\n"
            "- Redis, Kafka or similar message infrastructure\n"
            "- Experience with system design at scale"
        ),
        "required": ["Python", "FastAPI", "PostgreSQL", "REST API", "SQL", "Docker"],
        "preferred": ["Kubernetes", "AWS", "Redis", "Kafka"],
        "education": ["Bachelor's degree"],
        "responsibilities": [
            "Design and build REST APIs using Python and FastAPI",
            "Model and optimise PostgreSQL schemas and queries",
            "Implement background processing with Redis and task queues",
            "Write automated tests and maintain CI pipelines",
            "Participate in on-call and incident response",
        ],
    },
    {
        "title": "Data Analyst",
        "department": "Data",
        "location": "Pune, Maharashtra",
        "work_mode": WorkMode.ONSITE,
        "employment_type": EmploymentType.FULL_TIME,
        "min_exp": 2, "max_exp": 5,
        "salary": (1200000, 2000000),
        "openings": 2,
        "status": JobStatus.PUBLISHED,
        "description": (
            "Help our commercial teams make better decisions with data.\n\n"
            "Responsibilities:\n"
            "- Build dashboards and reports in Power BI and Tableau\n"
            "- Write complex SQL queries against our data warehouse\n"
            "- Analyse product and revenue trends and present findings\n"
            "- Partner with engineering on data quality\n\n"
            "Requirements:\n"
            "- 2+ years in an analytics role\n"
            "- Advanced SQL and Excel\n"
            "- Experience with Power BI or Tableau\n"
            "- Python for data analysis (pandas)\n"
            "- Bachelor's degree in Statistics, Mathematics, Economics or similar\n\n"
            "Preferred:\n"
            "- Snowflake or BigQuery\n"
            "- dbt and data modelling experience"
        ),
        "required": ["SQL", "Excel", "Power BI", "Python", "Pandas"],
        "preferred": ["Snowflake", "dbt", "Tableau"],
        "education": ["Bachelor's degree"],
        "responsibilities": [
            "Build dashboards and reports in Power BI and Tableau",
            "Write complex SQL queries against our data warehouse",
            "Analyse product and revenue trends and present findings",
            "Partner with engineering on data quality",
        ],
    },
    {
        "title": "Product Designer (UI/UX)",
        "department": "Design",
        "location": "Remote - India",
        "work_mode": WorkMode.REMOTE,
        "employment_type": EmploymentType.FULL_TIME,
        "min_exp": 3, "max_exp": 6,
        "salary": (1600000, 2600000),
        "openings": 1,
        "status": JobStatus.PUBLISHED,
        "description": (
            "Own the end-to-end design of features across our product.\n\n"
            "Responsibilities:\n"
            "- Design user flows, wireframes and high-fidelity interfaces in Figma\n"
            "- Conduct user research and usability testing\n"
            "- Maintain and evolve our design system\n"
            "- Work closely with engineers through implementation\n\n"
            "Requirements:\n"
            "- 3+ years designing digital products\n"
            "- Expert with Figma and prototyping tools\n"
            "- Strong portfolio demonstrating UX thinking\n"
            "- Experience with design systems and accessibility (WCAG)\n\n"
            "Preferred:\n"
            "- Familiarity with HTML and CSS\n"
            "- Motion design experience"
        ),
        "required": ["Figma", "User Experience Design", "Design Systems", "Prototyping",
                     "User Research"],
        "preferred": ["HTML", "CSS", "Accessibility"],
        "education": ["Bachelor's degree"],
        "responsibilities": [
            "Design user flows, wireframes and high-fidelity interfaces in Figma",
            "Conduct user research and usability testing",
            "Maintain and evolve our design system",
            "Work closely with engineers through implementation",
        ],
    },
    {
        "title": "DevOps Engineer",
        "department": "Engineering",
        "location": "Bengaluru, Karnataka",
        "work_mode": WorkMode.HYBRID,
        "employment_type": EmploymentType.FULL_TIME,
        "min_exp": 4, "max_exp": 9,
        "salary": (2400000, 3800000),
        "openings": 1,
        "status": JobStatus.PUBLISHED,
        "description": (
            "Own the infrastructure our engineering teams build on.\n\n"
            "Responsibilities:\n"
            "- Manage Kubernetes clusters across multiple environments\n"
            "- Build and maintain CI/CD pipelines\n"
            "- Automate infrastructure with Terraform\n"
            "- Improve observability with Prometheus and Grafana\n"
            "- Drive incident response and reliability improvements\n\n"
            "Requirements:\n"
            "- 4+ years in DevOps, SRE or platform engineering\n"
            "- Strong AWS and Kubernetes experience\n"
            "- Terraform and infrastructure-as-code\n"
            "- Docker and CI/CD tooling\n"
            "- Linux administration and shell scripting"
        ),
        "required": ["AWS", "Kubernetes", "Terraform", "Docker", "CI/CD", "Linux"],
        "preferred": ["Prometheus", "Grafana", "Ansible", "Python"],
        "education": ["Bachelor's degree"],
        "responsibilities": [
            "Manage Kubernetes clusters across multiple environments",
            "Build and maintain CI/CD pipelines",
            "Automate infrastructure with Terraform",
            "Improve observability with Prometheus and Grafana",
            "Drive incident response and reliability improvements",
        ],
    },
    {
        "title": "Frontend Developer (Fresher)",
        "department": "Engineering",
        "location": "Bengaluru, Karnataka",
        "work_mode": WorkMode.ONSITE,
        "employment_type": EmploymentType.FRESHER,
        "min_exp": 0, "max_exp": 2,
        "salary": (600000, 1000000),
        "openings": 4,
        "status": JobStatus.PUBLISHED,
        "description": (
            "A graduate role for engineers starting their career in frontend development.\n\n"
            "Responsibilities:\n"
            "- Build UI components under the guidance of senior engineers\n"
            "- Fix bugs and write tests\n"
            "- Learn our codebase, tooling and review process\n\n"
            "Requirements:\n"
            "- 0-2 years of experience; recent graduates welcome\n"
            "- Working knowledge of JavaScript, HTML and CSS\n"
            "- Familiarity with React from projects or internships\n"
            "- Bachelor's degree in Computer Science or related field"
        ),
        "required": ["JavaScript", "HTML", "CSS", "React"],
        "preferred": ["Git", "TypeScript"],
        "education": ["Bachelor's degree"],
        "responsibilities": [
            "Build UI components under the guidance of senior engineers",
            "Fix bugs and write tests",
            "Learn our codebase, tooling and review process",
        ],
    },
    {
        "title": "QA Automation Engineer",
        "department": "Engineering",
        "location": "Pune, Maharashtra",
        "work_mode": WorkMode.HYBRID,
        "employment_type": EmploymentType.FULL_TIME,
        "min_exp": 3, "max_exp": 6,
        "salary": (1400000, 2200000),
        "openings": 1,
        "status": JobStatus.PUBLISHED,
        "description": (
            "Build the automated testing that keeps our releases safe.\n\n"
            "Responsibilities:\n"
            "- Design and maintain end-to-end test suites with Playwright and Cypress\n"
            "- Build API test coverage with pytest\n"
            "- Integrate testing into CI/CD pipelines\n"
            "- Triage and report defects with clear reproduction steps\n\n"
            "Requirements:\n"
            "- 3+ years in QA automation\n"
            "- Strong Selenium, Cypress or Playwright experience\n"
            "- Python or JavaScript programming ability\n"
            "- Experience with CI/CD and Git"
        ),
        "required": ["Selenium", "Cypress", "Python", "QA Automation", "CI/CD"],
        "preferred": ["Playwright", "pytest", "Docker"],
        "education": ["Bachelor's degree"],
        "responsibilities": [
            "Design and maintain end-to-end test suites with Playwright and Cypress",
            "Build API test coverage with pytest",
            "Integrate testing into CI/CD pipelines",
            "Triage and report defects with clear reproduction steps",
        ],
    },
    {
        "title": "Engineering Manager",
        "department": "Engineering",
        "location": "Bengaluru, Karnataka",
        "work_mode": WorkMode.HYBRID,
        "employment_type": EmploymentType.FULL_TIME,
        "min_exp": 8, "max_exp": 14,
        "salary": (4000000, 6000000),
        "openings": 1,
        "status": JobStatus.PUBLISHED,
        "description": (
            "Lead one of our product engineering teams.\n\n"
            "Responsibilities:\n"
            "- Manage and grow a team of 6-10 engineers\n"
            "- Own delivery, quality and technical direction for your area\n"
            "- Partner with product on roadmap and prioritisation\n"
            "- Coach engineers through career development\n\n"
            "Requirements:\n"
            "- 8+ years in software engineering, including 2+ managing engineers\n"
            "- Strong background in system design and modern web architecture\n"
            "- Demonstrated experience hiring and developing engineers\n"
            "- Bachelor's degree in Computer Science or equivalent"
        ),
        "required": ["Leadership", "System Design", "Agile", "People Management"],
        "preferred": ["Python", "React", "AWS"],
        "education": ["Bachelor's degree"],
        "responsibilities": [
            "Manage and grow a team of 6-10 engineers",
            "Own delivery, quality and technical direction for your area",
            "Partner with product on roadmap and prioritisation",
            "Coach engineers through career development",
        ],
    },
    {
        "title": "Product Manager",
        "department": "Product",
        "location": "Bengaluru, Karnataka",
        "work_mode": WorkMode.HYBRID,
        "employment_type": EmploymentType.FULL_TIME,
        "min_exp": 4, "max_exp": 8,
        "salary": (2600000, 4000000),
        "openings": 1,
        "status": JobStatus.DRAFT,
        "description": (
            "Own a core area of the Northwind product from discovery to launch.\n\n"
            "Responsibilities:\n"
            "- Define product strategy and roadmap for your area\n"
            "- Run discovery with customers and translate insight into requirements\n"
            "- Work with engineering and design through delivery\n"
            "- Define and track success metrics\n\n"
            "Requirements:\n"
            "- 4+ years in product management for a software product\n"
            "- Strong analytical skills and comfort with data\n"
            "- Excellent written and verbal communication\n"
            "- Experience with roadmapping and user research"
        ),
        "required": ["Product Management", "Roadmapping", "User Research", "Communication"],
        "preferred": ["SQL", "Agile"],
        "education": ["Bachelor's degree"],
        "responsibilities": [
            "Define product strategy and roadmap for your area",
            "Run discovery with customers and translate insight into requirements",
            "Work with engineering and design through delivery",
            "Define and track success metrics",
        ],
    },
    {
        "title": "Software Engineering Intern",
        "department": "Engineering",
        "location": "Bengaluru, Karnataka",
        "work_mode": WorkMode.ONSITE,
        "employment_type": EmploymentType.INTERNSHIP,
        "min_exp": 0, "max_exp": 1,
        "salary": (300000, 480000),
        "openings": 6,
        "status": JobStatus.PUBLISHED,
        "description": (
            "A six-month internship for students and recent graduates.\n\n"
            "Responsibilities:\n"
            "- Work on a real product feature with a mentor\n"
            "- Write code, tests and documentation\n"
            "- Present your project at the end of the internship\n\n"
            "Requirements:\n"
            "- Currently studying or recently graduated in Computer Science\n"
            "- Programming ability in Python, JavaScript or Java\n"
            "- Understanding of data structures and algorithms\n"
            "- Curiosity and willingness to learn"
        ),
        "required": ["Python", "JavaScript", "Git"],
        "preferred": ["React", "SQL"],
        "education": ["Bachelor's degree"],
        "responsibilities": [
            "Work on a real product feature with a mentor",
            "Write code, tests and documentation",
            "Present your project at the end of the internship",
        ],
    },
]

# ----------------------------------------------------------------- candidates
CANDIDATE_TEMPLATES: list[dict] = [
    # --- strong React candidates
    {
        "first": "Rahul", "last": "Sharma", "city": "Bengaluru",
        "designation": "Senior Frontend Engineer", "company": "Flipkart",
        "years": 6.0, "notice": 30, "expected": 3000000,
        "skills": ["React", "TypeScript", "JavaScript", "REST API", "Git", "CSS",
                   "Next.js", "Redux", "Jest", "HTML"],
        "degree": "B.Tech in Computer Science", "level": "BACHELORS",
        "institution": "NIT Trichy", "grad": 2018,
        "target": "Senior React Developer",
        "roles": [
            ("Flipkart", "Senior Frontend Engineer", 2021, None, [
                "Built and maintained reusable React components for the seller design system",
                "Developed REST API integrations for the catalogue service",
                "Improved Core Web Vitals, cutting LCP by 40 percent",
                "Mentored three junior developers and led frontend code reviews",
            ]),
            ("Zoho", "Frontend Engineer", 2018, 2021, [
                "Developed customer-facing React and TypeScript interfaces",
                "Collaborated with design on a shared component library",
            ]),
        ],
    },
    {
        "first": "Priya", "last": "Shah", "city": "Pune",
        "designation": "Frontend Lead", "company": "Razorpay",
        "years": 5.5, "notice": 60, "expected": 2900000,
        "skills": ["React", "TypeScript", "REST API", "Git", "JavaScript", "CSS",
                   "GraphQL", "Docker", "AWS"],
        "degree": "B.E. in Information Technology", "level": "BACHELORS",
        "institution": "COEP Pune", "grad": 2019,
        "target": "Senior React Developer",
        "roles": [
            ("Razorpay", "Frontend Lead", 2022, None, [
                "Led the frontend for the merchant dashboard used by 200k businesses",
                "Built reusable React components and owned the design system",
                "Integrated REST and GraphQL APIs across six product surfaces",
                "Mentored junior developers through structured code review",
            ]),
            ("Freshworks", "Software Engineer", 2019, 2022, [
                "Built React interfaces for the CRM product",
                "Wrote unit and integration tests with Jest",
            ]),
        ],
    },
    {
        "first": "Amit", "last": "Patel", "city": "Ahmedabad",
        "designation": "Software Engineer", "company": "Infosys",
        "years": 4.0, "notice": 90, "expected": 2200000,
        "skills": ["React", "JavaScript", "REST API", "Git", "HTML", "CSS", "Java"],
        "degree": "B.Tech in Computer Engineering", "level": "BACHELORS",
        "institution": "Nirma University", "grad": 2020,
        "target": "Senior React Developer",
        "roles": [
            ("Infosys", "Software Engineer", 2020, None, [
                "Developed React front ends for banking clients",
                "Consumed REST APIs and handled authentication flows",
                "Participated in code reviews and agile ceremonies",
            ]),
        ],
    },
    # --- backend candidates
    {
        "first": "Neha", "last": "Joshi", "city": "Bengaluru",
        "designation": "Senior Backend Engineer", "company": "Swiggy",
        "years": 5.0, "notice": 45, "expected": 2900000,
        "skills": ["Python", "FastAPI", "PostgreSQL", "REST API", "Docker", "SQL",
                   "Redis", "Kubernetes", "AWS", "Kafka"],
        "degree": "M.Tech in Computer Science", "level": "MASTERS",
        "institution": "IIIT Hyderabad", "grad": 2019,
        "target": "Backend Engineer (Python)",
        "roles": [
            ("Swiggy", "Senior Backend Engineer", 2021, None, [
                "Designed and built REST APIs with Python and FastAPI",
                "Optimised PostgreSQL schemas and queries for the order service",
                "Implemented background processing with Redis and Celery",
                "Participated in on-call rotation and incident response",
            ]),
            ("Practo", "Backend Engineer", 2019, 2021, [
                "Built Django services and REST APIs",
                "Wrote automated tests and maintained CI pipelines",
            ]),
        ],
    },
    {
        "first": "Rohan", "last": "Mehta", "city": "Hyderabad",
        "designation": "Backend Engineer", "company": "PhonePe",
        "years": 3.5, "notice": 30, "expected": 2400000,
        "skills": ["Python", "Django", "PostgreSQL", "REST API", "SQL", "Docker", "Git"],
        "degree": "B.Tech in Computer Science", "level": "BACHELORS",
        "institution": "VIT Vellore", "grad": 2021,
        "target": "Backend Engineer (Python)",
        "roles": [
            ("PhonePe", "Backend Engineer", 2021, None, [
                "Built REST APIs in Python for the payments platform",
                "Modelled PostgreSQL schemas and tuned slow queries",
                "Containerised services with Docker",
            ]),
        ],
    },
    {
        "first": "Sanjay", "last": "Kulkarni", "city": "Pune",
        "designation": "Software Engineer", "company": "TCS",
        "years": 3.0, "notice": 90, "expected": 1800000,
        "skills": ["Java", "Spring Boot", "SQL", "REST API", "Git"],
        "degree": "B.E. in Computer Science", "level": "BACHELORS",
        "institution": "Pune University", "grad": 2021,
        "target": "Backend Engineer (Python)",
        "roles": [
            ("TCS", "Software Engineer", 2021, None, [
                "Developed Spring Boot microservices for an insurance client",
                "Wrote SQL queries and stored procedures",
            ]),
        ],
    },
    # --- data candidates
    {
        "first": "Ananya", "last": "Reddy", "city": "Pune",
        "designation": "Data Analyst", "company": "Deloitte",
        "years": 3.0, "notice": 60, "expected": 1700000,
        "skills": ["SQL", "Excel", "Power BI", "Python", "Pandas", "Tableau", "Snowflake"],
        "degree": "M.Sc in Statistics", "level": "MASTERS",
        "institution": "University of Pune", "grad": 2021,
        "target": "Data Analyst",
        "roles": [
            ("Deloitte", "Data Analyst", 2021, None, [
                "Built dashboards and reports in Power BI and Tableau",
                "Wrote complex SQL queries against the data warehouse",
                "Analysed revenue trends and presented findings to leadership",
                "Partnered with engineering on data quality initiatives",
            ]),
        ],
    },
    {
        "first": "Manish", "last": "Verma", "city": "Mumbai",
        "designation": "Business Analyst", "company": "Accenture",
        "years": 2.5, "notice": 30, "expected": 1400000,
        "skills": ["SQL", "Excel", "Power BI", "Python"],
        "degree": "MBA in Analytics", "level": "MASTERS",
        "institution": "NMIMS Mumbai", "grad": 2022,
        "target": "Data Analyst",
        "roles": [
            ("Accenture", "Business Analyst", 2022, None, [
                "Created Power BI dashboards for client reporting",
                "Wrote SQL queries to extract and validate data",
            ]),
        ],
    },
    # --- design candidates
    {
        "first": "Kavya", "last": "Nambiar", "city": "Kochi",
        "designation": "Senior Product Designer", "company": "Zoho",
        "years": 5.0, "notice": 30, "expected": 2400000,
        "skills": ["Figma", "User Experience Design", "Design Systems", "Prototyping",
                   "User Research", "HTML", "CSS", "Accessibility"],
        "degree": "B.Des in Interaction Design", "level": "BACHELORS",
        "institution": "NID Ahmedabad", "grad": 2019,
        "target": "Product Designer (UI/UX)",
        "roles": [
            ("Zoho", "Senior Product Designer", 2021, None, [
                "Designed user flows, wireframes and high-fidelity interfaces in Figma",
                "Conducted user research and usability testing with 60+ participants",
                "Maintained and evolved the product design system",
                "Worked closely with engineers through implementation",
            ]),
            ("Freelance", "Product Designer", 2019, 2021, [
                "Designed mobile and web products for early-stage startups",
            ]),
        ],
    },
    # --- devops candidates
    {
        "first": "Arun", "last": "Pillai", "city": "Bengaluru",
        "designation": "DevOps Engineer", "company": "Nutanix",
        "years": 6.0, "notice": 60, "expected": 3400000,
        "skills": ["AWS", "Kubernetes", "Terraform", "Docker", "CI/CD", "Linux",
                   "Prometheus", "Grafana", "Ansible", "Python"],
        "degree": "B.Tech in Electronics", "level": "BACHELORS",
        "institution": "CUSAT", "grad": 2018,
        "target": "DevOps Engineer",
        "roles": [
            ("Nutanix", "DevOps Engineer", 2020, None, [
                "Managed Kubernetes clusters across staging and production",
                "Built and maintained CI/CD pipelines with Jenkins and GitHub Actions",
                "Automated infrastructure with Terraform across three AWS regions",
                "Improved observability with Prometheus and Grafana",
                "Drove incident response and post-mortem culture",
            ]),
            ("Wipro", "Systems Engineer", 2018, 2020, [
                "Administered Linux servers and automated deployments with Ansible",
            ]),
        ],
    },
    # --- freshers
    {
        "first": "Sneha", "last": "Bhat", "city": "Bengaluru",
        "designation": None, "company": None,
        "years": 0.5, "notice": 0, "expected": 800000,
        "skills": ["JavaScript", "HTML", "CSS", "React", "Git"],
        "degree": "B.Tech in Computer Science", "level": "BACHELORS",
        "institution": "PES University", "grad": 2024,
        "target": "Frontend Developer (Fresher)",
        "roles": [
            ("Zeta", "Frontend Intern", 2024, None, [
                "Built UI components in React under senior guidance",
                "Fixed bugs and wrote unit tests",
            ]),
        ],
    },
    {
        "first": "Aditya", "last": "Rane", "city": "Mumbai",
        "designation": None, "company": None,
        "years": 0.0, "notice": 0, "expected": 700000,
        "skills": ["JavaScript", "HTML", "CSS", "Python"],
        "degree": "B.E. in Information Technology", "level": "BACHELORS",
        "institution": "Mumbai University", "grad": 2024,
        "target": "Frontend Developer (Fresher)",
        "roles": [],
    },
    {
        "first": "Ishita", "last": "Agarwal", "city": "Delhi",
        "designation": None, "company": None,
        "years": 0.0, "notice": 0, "expected": 400000,
        "skills": ["Python", "JavaScript", "Git", "SQL"],
        "degree": "B.Tech in Computer Science", "level": "BACHELORS",
        "institution": "Delhi Technological University", "grad": 2025,
        "target": "Software Engineering Intern",
        "roles": [],
    },
    {
        "first": "Varun", "last": "Nayak", "city": "Bengaluru",
        "designation": None, "company": None,
        "years": 0.0, "notice": 0, "expected": 420000,
        "skills": ["Java", "Python", "Git"],
        "degree": "B.Tech in Information Science", "level": "BACHELORS",
        "institution": "RV College of Engineering", "grad": 2025,
        "target": "Software Engineering Intern",
        "roles": [],
    },
    # --- QA
    {
        "first": "Deepa", "last": "Chandran", "city": "Pune",
        "designation": "QA Automation Engineer", "company": "Persistent Systems",
        "years": 4.0, "notice": 60, "expected": 1900000,
        "skills": ["Selenium", "Cypress", "Python", "QA Automation", "CI/CD",
                   "Playwright", "pytest", "Git"],
        "degree": "B.E. in Computer Science", "level": "BACHELORS",
        "institution": "Savitribai Phule Pune University", "grad": 2020,
        "target": "QA Automation Engineer",
        "roles": [
            ("Persistent Systems", "QA Automation Engineer", 2020, None, [
                "Designed and maintained end-to-end suites with Playwright and Cypress",
                "Built API test coverage with pytest",
                "Integrated testing into CI/CD pipelines",
                "Triaged defects with clear reproduction steps",
            ]),
        ],
    },
    # --- weaker / mismatched candidates, so the ATS distribution is realistic
    {
        "first": "Suresh", "last": "Babu", "city": "Chennai",
        "designation": "PHP Developer", "company": "Local Agency",
        "years": 7.0, "notice": 15, "expected": 1500000,
        "skills": ["PHP", "WordPress", "MySQL", "jQuery"],
        "degree": "B.Sc in Computer Science", "level": "BACHELORS",
        "institution": "Madras University", "grad": 2016,
        "target": "Senior React Developer",
        "roles": [
            ("Local Agency", "PHP Developer", 2016, None, [
                "Built WordPress websites for small business clients",
                "Maintained legacy PHP applications",
            ]),
        ],
    },
    {
        "first": "Pooja", "last": "Malhotra", "city": "Gurugram",
        "designation": "Support Engineer", "company": "HCL",
        "years": 3.0, "notice": 30, "expected": 1000000,
        "skills": ["Excel", "Communication", "SQL"],
        "degree": "B.Com", "level": "BACHELORS",
        "institution": "Delhi University", "grad": 2021,
        "target": "Data Analyst",
        "roles": [
            ("HCL", "Support Engineer", 2021, None, [
                "Handled customer tickets and escalations",
                "Produced weekly Excel reports on ticket volumes",
            ]),
        ],
    },
    {
        "first": "Nikhil", "last": "Saxena", "city": "Noida",
        "designation": "Full Stack Developer", "company": "Paytm",
        "years": 4.5, "notice": 45, "expected": 2600000,
        "skills": ["React", "Node.js", "JavaScript", "MongoDB", "REST API", "Git",
                   "TypeScript", "Docker"],
        "degree": "B.Tech in Computer Science", "level": "BACHELORS",
        "institution": "Amity University", "grad": 2020,
        "target": "Senior React Developer",
        "roles": [
            ("Paytm", "Full Stack Developer", 2020, None, [
                "Built React front ends and Node.js REST APIs",
                "Developed reusable component libraries",
                "Integrated third-party payment APIs",
            ]),
        ],
    },
    {
        "first": "Lakshmi", "last": "Narayan", "city": "Bengaluru",
        "designation": "Engineering Manager", "company": "Myntra",
        "years": 11.0, "notice": 90, "expected": 5500000,
        "skills": ["Leadership", "System Design", "Agile", "People Management",
                   "Python", "React", "AWS"],
        "degree": "M.Tech in Computer Science", "level": "MASTERS",
        "institution": "IIT Madras", "grad": 2014,
        "target": "Engineering Manager",
        "roles": [
            ("Myntra", "Engineering Manager", 2020, None, [
                "Managed and grew a team of nine engineers",
                "Owned delivery, quality and technical direction for the checkout area",
                "Partnered with product on roadmap and prioritisation",
                "Coached engineers through structured career development",
            ]),
            ("Amazon", "Senior Software Engineer", 2016, 2020, [
                "Designed distributed systems for the fulfilment platform",
            ]),
            ("Oracle", "Software Engineer", 2014, 2016, [
                "Built backend services in Java",
            ]),
        ],
    },
    {
        "first": "Tanvi", "last": "Deshmukh", "city": "Remote",
        "designation": "UI Designer", "company": "Startup",
        "years": 2.0, "notice": 15, "expected": 1200000,
        "skills": ["Figma", "User Interface Design", "Prototyping", "Adobe XD"],
        "degree": "B.Des", "level": "BACHELORS",
        "institution": "MIT Institute of Design", "grad": 2022,
        "target": "Product Designer (UI/UX)",
        "roles": [
            ("Startup", "UI Designer", 2022, None, [
                "Designed mobile app interfaces in Figma",
                "Created prototypes for user testing",
            ]),
        ],
    },
]

SOURCES = [
    ApplicationSource.LINKEDIN,
    ApplicationSource.COMPANY_WEBSITE,
    ApplicationSource.JOB_BOARD,
    ApplicationSource.REFERRAL,
    ApplicationSource.INSTAGRAM,
    ApplicationSource.DIRECT,
]


def _months_before(anchor: date, months: int) -> date:
    """``anchor`` shifted back by N months, clamped to the first of the month."""
    total = anchor.year * 12 + (anchor.month - 1) - months
    return date(total // 12, total % 12 + 1, 1)


def build_resume_docx(text: str) -> bytes:
    """Render resume text into a real .docx.

    The seed deliberately produces a genuine Office document rather than a text blob, so
    the upload path, the DOCX extractor and the parser are all exercised exactly as they
    would be by a real candidate's file.
    """
    from io import BytesIO

    import docx

    document = docx.Document()
    for line in text.split("\n"):
        if not line.strip():
            document.add_paragraph("")
        elif line.isupper() and len(line) < 30:
            document.add_heading(line.title(), level=2)
        else:
            document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_resume_text(template: dict) -> str:
    """Compose a realistic plain-text resume the real parser will process."""
    lines: list[str] = [
        f"{template['first']} {template['last']}",
        f"{template['first'].lower()}.{template['last'].lower()}@example.test",
        f"+91 98{RANDOM.randint(10000000, 99999999)}",
        f"{template['city']}, India",
        f"https://linkedin.com/in/{template['first'].lower()}{template['last'].lower()}",
        f"https://github.com/{template['first'].lower()}{template['last'].lower()}",
        "",
        "SUMMARY",
    ]
    if template["designation"]:
        lines.append(
            f"{template['designation']} with {template['years']:g} years of experience "
            f"building production software. Currently at {template['company']}."
        )
    else:
        lines.append(
            "Recent graduate seeking a role where I can build production software and "
            "learn from experienced engineers."
        )

    lines += ["", "SKILLS", ", ".join(template["skills"]), "", "EXPERIENCE"]

    # Anchor the employment history to the declared total experience rather than to
    # hard-coded years, so the resume the parser reads always agrees with the profile.
    # Without this the seed drifts out of sync as real-world time passes.
    roles = template["roles"]
    if roles:
        # Split the career into contiguous, gap-free spans, newest first, so the dates on
        # the resume sum to exactly the declared total experience. Month-level precision
        # keeps the parser's computed total within a rounding error of the declaration,
        # and exercises the month-name date path.
        total_months = max(1, round(template["years"] * 12))
        base, remainder = divmod(total_months, len(roles))
        spans = [max(1, base + (1 if i < remainder else 0)) for i in range(len(roles))]

        cursor_end = date.today().replace(day=1)  # walk backwards from this month
        for index, (company, position, _s, _e, bullets) in enumerate(roles):
            start = _months_before(cursor_end, spans[index])
            end_label = "Present" if index == 0 else cursor_end.strftime("%b %Y")
            lines.append(
                f"{position} at {company}    {start.strftime('%b %Y')} - {end_label}"
            )
            lines += [f"- {b}" for b in bullets]
            lines.append("")
            cursor_end = start
    else:
        lines.append("No professional experience yet - recent graduate.")
        lines.append("")

    lines += [
        "EDUCATION",
        f"{template['degree']} from {template['institution']}    "
        f"{template['grad'] - 4} - {template['grad']}",
        "",
        "PROJECTS",
        "- Built a personal portfolio site and several open-source utilities",
    ]
    return "\n".join(lines)


async def _reset(session: AsyncSession) -> None:
    """Delete demo data. Deliberately scoped to the demo company plus seeded users."""

    company = await session.scalar(select(Company).where(Company.name == COMPANY_NAME))
    if company is not None:
        await session.execute(delete(Company).where(Company.id == company.id))

    seed_emails = [s["email"] for s in STAFF] + [settings.SEED_SUPER_ADMIN_EMAIL]
    await session.execute(delete(User).where(User.email.in_(seed_emails)))
    # Candidate logins created by the seed all share this domain.
    await session.execute(delete(User).where(User.email.like("%@example.test")))
    await session.flush()
    logger.info("seed_reset_complete")


async def seed(*, reset: bool = False) -> dict:
    configure_logging()
    await ensure_schema()

    async with session_scope() as session:
        await bootstrap_database(session)

    if reset:
        async with session_scope() as session:
            await _reset(session)

    summary: dict = {}

    # ------------------------------------------------------------ super admin
    async with session_scope() as session:
        existing = await session.scalar(
            select(User).where(User.email == settings.SEED_SUPER_ADMIN_EMAIL)
        )
        if existing is None:
            role = await session.scalar(
                select(Role).where(
                    Role.name == RoleName.SUPER_ADMIN.value, Role.company_id.is_(None)
                )
            )
            admin = User(
                email=settings.SEED_SUPER_ADMIN_EMAIL,
                hashed_password=hash_password(settings.SEED_SUPER_ADMIN_PASSWORD),
                first_name="Platform",
                last_name="Admin",
                status=UserStatus.ACTIVE,
                email_verified_at=datetime.now(UTC),
            )
            admin.roles.append(role)
            session.add(admin)
            await session.flush()
            logger.info("super_admin_created", email=settings.SEED_SUPER_ADMIN_EMAIL)

    # ---------------------------------------------------------------- company
    async with session_scope() as session:
        company = await session.scalar(select(Company).where(Company.name == COMPANY_NAME))
        if company is not None:
            logger.info("seed_already_present", company=COMPANY_NAME)
            return {"status": "already_seeded", "company_id": str(company.id)}

        company = Company(
            name=COMPANY_NAME,
            slug=slugify(COMPANY_NAME),
            legal_name="Northwind Technologies Private Limited",
            description=(
                "Northwind Technologies builds commerce infrastructure used by thousands "
                "of businesses across India. We are a product company of about 400 "
                "people, headquartered in Bengaluru."
            ),
            website="https://northwind.test",
            industry="Software & Technology",
            size=CompanySize.SIZE_201_500,
            founded_year=2016,
            headquarters="Bengaluru, Karnataka, India",
            contact_email="talent@northwind.test",
            contact_phone="+91 8012345678",
            status=CompanyStatus.ACTIVE,
            subscription_plan="enterprise",
            settings={
                "interview_reminder_offsets_minutes": [1440, 60],
                "default_offer_validity_days": 7,
                "auto_score_on_apply": True,
                "require_offer_approval": True,
                "default_currency": "INR",
            },
        )
        session.add(company)
        await session.flush()
        company_id = company.id

        departments: dict[str, Department] = {}
        for name in DEPARTMENTS:
            department = Department(company_id=company_id, name=name)
            session.add(department)
            departments[name] = department

        for location in LOCATIONS:
            session.add(CompanyLocation(company_id=company_id, **location))
        await session.flush()

        # ------------------------------------------------------------- staff
        roles_by_name = {
            role.name: role
            for role in (
                await session.execute(select(Role).where(Role.company_id.is_(None)))
            )
            .scalars()
            .all()
        }
        staff_users: dict[str, User] = {}
        for record in STAFF:
            user = User(
                email=record["email"],
                hashed_password=hash_password(settings.SEED_DEMO_PASSWORD),
                first_name=record["first_name"],
                last_name=record["last_name"],
                job_title=record["job_title"],
                company_id=company_id,
                status=UserStatus.ACTIVE,
                email_verified_at=datetime.now(UTC),
                timezone="Asia/Kolkata",
            )
            for role_name in record["roles"]:
                user.roles.append(roles_by_name[role_name.value])
            session.add(user)
            staff_users[record["email"]] = user
        await session.flush()

        # ------------------------------- defaults: ATS profile & email templates
        from app.modules.ats.service import AtsService
        from app.modules.emails.service import EmailService

        await AtsService(session, company_id).ensure_default_profile()
        await EmailService(session, company_id).ensure_default_templates()

        recruiters = [
            staff_users["priya.nair@northwind.test"],
            staff_users["arjun.mehta@northwind.test"],
            staff_users["sneha.rao@northwind.test"],
        ]
        managers = [
            staff_users["vikram.desai@northwind.test"],
            staff_users["meera.krishnan@northwind.test"],
        ]
        interviewers = [
            staff_users["rohan.gupta@northwind.test"],
            staff_users["anita.shah@northwind.test"],
            staff_users["karthik.iyer@northwind.test"],
            staff_users["divya.menon@northwind.test"],
            staff_users["sameer.khan@northwind.test"],
        ]

        # -------------------------------------------------------------- jobs
        from app.utils.skills import categorise_skill, display_skill, normalise_skill

        jobs_by_title: dict[str, Job] = {}
        for index, spec in enumerate(JOBS):
            recruiter = recruiters[index % len(recruiters)]
            manager = managers[index % len(managers)]
            published_at = (
                datetime.now(UTC) - timedelta(days=RANDOM.randint(5, 45))
                if spec["status"] == JobStatus.PUBLISHED
                else None
            )

            job = Job(
                company_id=company_id,
                title=spec["title"],
                slug=slugify(spec["title"]),
                reference_code=f"{spec['department'][:3].upper()}-2024-{1000 + index}",
                description=spec["description"],
                department_id=departments[spec["department"]].id,
                location_text=spec["location"],
                work_mode=spec["work_mode"],
                employment_type=spec["employment_type"],
                min_experience_years=spec["min_exp"],
                max_experience_years=spec["max_exp"],
                salary_min=spec["salary"][0],
                salary_max=spec["salary"][1],
                salary_currency="INR",
                openings=spec["openings"],
                responsibilities=spec["responsibilities"],
                education_requirements=spec["education"],
                benefits=[
                    "Health insurance for you and your family",
                    "Annual learning budget",
                    "Flexible working hours",
                    "Employee stock options",
                ],
                status=spec["status"],
                published_at=published_at,
                created_by_id=recruiter.id,
                hiring_manager_id=manager.id,
                application_deadline=date.today() + timedelta(days=RANDOM.randint(20, 90)),
            )
            session.add(job)
            await session.flush()

            for name in spec["required"]:
                session.add(
                    JobSkill(
                        job_id=job.id,
                        name=display_skill(name),
                        normalised_name=normalise_skill(name),
                        importance=SkillImportance.REQUIRED,
                        weight=5 if name in spec["required"][:2] else 3,
                        category=categorise_skill(name),
                        source="MANUAL",
                    )
                )
            for name in spec["preferred"]:
                session.add(
                    JobSkill(
                        job_id=job.id,
                        name=display_skill(name),
                        normalised_name=normalise_skill(name),
                        importance=SkillImportance.PREFERRED,
                        weight=2,
                        category=categorise_skill(name),
                        source="MANUAL",
                    )
                )

            session.add(
                JobScreeningQuestion(
                    job_id=job.id,
                    question="How many years of relevant professional experience do you have?",
                    question_type=ScreeningQuestionType.EXPERIENCE,
                    is_required=True,
                    display_order=0,
                    scoring={"min": spec["min_exp"], "points": 20},
                )
            )
            session.add(
                JobScreeningQuestion(
                    job_id=job.id,
                    question="What is your notice period in days?",
                    question_type=ScreeningQuestionType.NOTICE_PERIOD,
                    is_required=True,
                    display_order=1,
                    scoring={"max": 60, "points": 10},
                )
            )
            session.add(
                JobScreeningQuestion(
                    job_id=job.id,
                    question=f"Are you able to work from {spec['location']}?",
                    question_type=ScreeningQuestionType.YES_NO,
                    is_required=True,
                    display_order=2,
                    scoring={"expected": "YES", "points": 10},
                    is_knockout=True,
                )
            )
            session.add(
                JobScreeningQuestion(
                    job_id=job.id,
                    question="Briefly describe a project you are proud of.",
                    question_type=ScreeningQuestionType.TEXT,
                    is_required=False,
                    display_order=3,
                )
            )

            session.add(
                JobHiringTeamMember(job_id=job.id, user_id=recruiter.id, team_role="RECRUITER")
            )
            session.add(
                JobHiringTeamMember(job_id=job.id, user_id=manager.id, team_role="MANAGER")
            )
            for interviewer in RANDOM.sample(interviewers, 2):
                session.add(
                    JobHiringTeamMember(
                        job_id=job.id, user_id=interviewer.id, team_role="INTERVIEWER"
                    )
                )

            jobs_by_title[spec["title"]] = job

        await session.flush()

        # ---------------------------------------------------------- workflows
        shortlist_workflow = Workflow(
            company_id=company_id,
            name="Auto-shortlist strong matches",
            description=(
                "When an ATS score of 82 or above is generated and at least 70% of the "
                "required skills are evidenced, move the application to Shortlisted and "
                "email the candidate."
            ),
            trigger=WorkflowTrigger.ATS_SCORE_GENERATED,
            conditions={
                "op": "AND",
                "rules": [
                    {"field": "ats_score", "operator": "gte", "value": 82},
                    {"field": "required_skills_match", "operator": "gte", "value": 70},
                    {"field": "application_status", "operator": "in",
                     "value": ["APPLIED", "UNDER_REVIEW"]},
                ],
            },
            is_enabled=True,
            requires_human_approval=False,
            priority=10,
            created_by_id=recruiters[0].id,
        )
        session.add(shortlist_workflow)
        await session.flush()
        session.add_all(
            [
                WorkflowStep(
                    workflow_id=shortlist_workflow.id,
                    step_order=0,
                    action_type=WorkflowActionType.CHANGE_STATUS,
                    config={"status": ApplicationStatus.SHORTLISTED.value},
                ),
                WorkflowStep(
                    workflow_id=shortlist_workflow.id,
                    step_order=1,
                    action_type=WorkflowActionType.SEND_EMAIL,
                    config={
                        "template_key": EmailTemplateKey.SHORTLISTED.value,
                        "custom_message": (
                            "Your background looks like a strong match for this role."
                        ),
                    },
                ),
                WorkflowStep(
                    workflow_id=shortlist_workflow.id,
                    step_order=2,
                    action_type=WorkflowActionType.NOTIFY,
                    config={
                        "title": "Strong match auto-shortlisted",
                        "message": "A candidate was shortlisted automatically. Review when you can.",
                    },
                ),
            ]
        )

        review_workflow = Workflow(
            company_id=company_id,
            name="Flag low scores for manual review",
            description=(
                "Applications scoring under 45 are flagged for a human to look at. They "
                "are never rejected automatically."
            ),
            trigger=WorkflowTrigger.ATS_SCORE_GENERATED,
            conditions={
                "op": "AND",
                "rules": [{"field": "ats_score", "operator": "lt", "value": 45}],
            },
            is_enabled=True,
            requires_human_approval=False,
            priority=20,
            created_by_id=recruiters[0].id,
        )
        session.add(review_workflow)
        await session.flush()
        session.add_all(
            [
                WorkflowStep(
                    workflow_id=review_workflow.id,
                    step_order=0,
                    action_type=WorkflowActionType.ADD_TAG,
                    config={"tag": "needs-manual-review"},
                ),
                WorkflowStep(
                    workflow_id=review_workflow.id,
                    step_order=1,
                    action_type=WorkflowActionType.FLAG_FOR_REVIEW,
                    config={
                        "code": "LOW_ATS_SCORE",
                        "message": (
                            "This application scored below the review threshold. Check "
                            "whether the resume parsed correctly before deciding."
                        ),
                    },
                ),
            ]
        )

        talent_pool_workflow = Workflow(
            company_id=company_id,
            name="Add rejected strong candidates to the talent pool",
            description=(
                "Keeps good candidates who were not selected available for future roles."
            ),
            trigger=WorkflowTrigger.APPLICATION_STATUS_CHANGED,
            conditions={
                "op": "AND",
                "rules": [
                    {"field": "new_status", "operator": "eq", "value": "REJECTED"},
                    {"field": "ats_score", "operator": "gte", "value": 70},
                ],
            },
            is_enabled=True,
            priority=30,
            created_by_id=recruiters[0].id,
        )
        session.add(talent_pool_workflow)
        await session.flush()
        session.add(
            WorkflowStep(
                workflow_id=talent_pool_workflow.id,
                step_order=0,
                action_type=WorkflowActionType.ADD_TO_TALENT_POOL,
                config={"pool_name": "Strong candidates - future roles"},
            )
        )
        await session.flush()

        summary["company_id"] = str(company_id)
        summary["jobs"] = len(jobs_by_title)
        summary["staff"] = len(staff_users)

    logger.info("seed_core_complete", **summary)

    # ------------------------------------------------- candidates & applications
    stats = await _seed_applications(uuid.UUID(summary["company_id"]))
    summary.update(stats)

    logger.info("seed_complete", **summary)
    return summary


async def _seed_applications(company_id: uuid.UUID) -> dict:
    """Create candidates and applications through the real intake + ATS pipeline."""
    from app.core.enums import RoleName as _RoleName
    from app.modules.applications.service import ApplicationIntakeService
    from app.modules.ats.service import AtsService
    from app.modules.resumes.service import ResumeService

    created_candidates = 0
    created_applications = 0
    scored = 0

    async with session_scope() as session:
        candidate_role = await session.scalar(
            select(Role).where(
                Role.name == _RoleName.CANDIDATE.value, Role.company_id.is_(None)
            )
        )
        jobs = {
            job.title: job
            for job in (
                await session.execute(
                    select(Job)
                    .where(Job.company_id == company_id)
                    .options(selectinload(Job.skills), selectinload(Job.screening_questions))
                )
            )
            .unique()
            .scalars()
            .all()
        }
        intake = ApplicationIntakeService(session, company_id)
        resume_service = ResumeService(session, company_id)

        published_titles = [t for t, j in jobs.items() if j.status == JobStatus.PUBLISHED]

        for template in CANDIDATE_TEMPLATES:
            email = f"{template['first'].lower()}.{template['last'].lower()}@example.test"

            # A login so the candidate portal is demonstrable.
            user = await session.scalar(select(User).where(User.email == email))
            if user is None:
                user = User(
                    email=email,
                    hashed_password=hash_password(settings.SEED_DEMO_PASSWORD),
                    first_name=template["first"],
                    last_name=template["last"],
                    status=UserStatus.ACTIVE,
                    email_verified_at=datetime.now(UTC),
                    timezone="Asia/Kolkata",
                )
                user.roles.append(candidate_role)
                session.add(user)
                await session.flush()

            candidate, created = await intake.find_or_create_candidate(
                email=email,
                first_name=template["first"],
                last_name=template["last"],
                phone=f"+9198{RANDOM.randint(10000000, 99999999)}",
                user_id=user.id,
                location=f"{template['city']}, India",
                source=ApplicationSource.LINKEDIN.value,
                extra={
                    "current_designation": template["designation"],
                    "current_company": template["company"],
                    "total_experience_years": template["years"],
                    "expected_salary": template["expected"],
                    "notice_period_days": template["notice"],
                    "email_verified": True,
                    "linkedin_url": (
                        f"https://linkedin.com/in/"
                        f"{template['first'].lower()}{template['last'].lower()}"
                    ),
                },
            )
            if created:
                created_candidates += 1

            # Store a real .docx and let the real extractor + parser handle it.
            resume = await resume_service.upload(
                candidate=candidate,
                filename=f"{template['first']}_{template['last']}_Resume.docx",
                content=build_resume_docx(build_resume_text(template)),
            )
            await session.flush()

            # Apply to the target job plus, sometimes, a second one.
            targets = [template["target"]]
            if RANDOM.random() < 0.35:
                other = RANDOM.choice(published_titles)
                if other != template["target"]:
                    targets.append(other)

            for target_title in targets:
                job = jobs.get(target_title)
                if job is None or job.status != JobStatus.PUBLISHED:
                    continue
                try:
                    application = await intake.create_application(
                        job=job,
                        candidate=candidate,
                        source=RANDOM.choice(SOURCES),
                        source_detail=None,
                        cover_letter=(
                            f"I am excited to apply for the {job.title} role at "
                            f"{COMPANY_NAME}. My background in "
                            f"{', '.join(template['skills'][:3])} lines up well with what "
                            "you are looking for."
                        ),
                        expected_salary=template["expected"],
                        notice_period_days=template["notice"],
                        consent_given=True,
                        actor_id=user.id,
                    )
                except Exception as exc:  # already applied, job closed, etc.
                    logger.debug("seed_application_skipped", reason=str(exc))
                    continue

                application.resume_id = resume.id
                # Backdate so the analytics charts have a spread of dates.
                days_ago = RANDOM.randint(1, 40)
                application.created_at = datetime.now(UTC) - timedelta(days=days_ago)
                application.status_changed_at = application.created_at
                created_applications += 1

            await session.flush()

        # The parser and scorer run against committed rows.
        await session.commit()

    # ------------------------------------------------------- parse + score
    async with session_scope() as session:
        from app.models.resume import Resume

        resume_ids = list(
            (
                await session.execute(
                    select(Resume.id).where(Resume.company_id == company_id)
                )
            )
            .scalars()
            .all()
        )

    for resume_id in resume_ids:
        async with session_scope() as session:
            try:
                await ResumeService(session, company_id).process(resume_id)
            except Exception as exc:
                logger.warning("seed_resume_parse_failed", resume_id=str(resume_id), error=str(exc))

    async with session_scope() as session:
        from app.models.application import Application

        application_ids = list(
            (
                await session.execute(
                    select(Application.id).where(Application.company_id == company_id)
                )
            )
            .scalars()
            .all()
        )

    for application_id in application_ids:
        async with session_scope() as session:
            try:
                await AtsService(session, company_id).score(application_id)
                scored += 1
            except Exception as exc:
                logger.warning(
                    "seed_scoring_failed", application_id=str(application_id), error=str(exc)
                )

    # ------------------------------------------- pipeline, interviews, offers
    await _seed_pipeline(company_id)

    return {
        "candidates": created_candidates,
        "applications": created_applications,
        "scored": scored,
    }


async def _seed_pipeline(company_id: uuid.UUID) -> None:
    """Move a realistic subset through screening, interviews, feedback and offers."""
    from app.models.application import Application
    from app.modules.applications.service import ApplicationPipelineService
    from app.modules.interviews.service import InterviewService
    from app.modules.offers.service import OfferService

    async with session_scope() as session:
        applications = list(
            (
                await session.execute(
                    select(Application)
                    .where(Application.company_id == company_id)
                    .options(
                        selectinload(Application.candidate), selectinload(Application.job)
                    )
                    .order_by(Application.ats_score.desc().nullslast())
                )
            )
            .unique()
            .scalars()
            .all()
        )

        pipeline = ApplicationPipelineService(session, company_id)
        recruiter = await session.scalar(
            select(User).where(User.email == "priya.nair@northwind.test")
        )

        # Spread applications across the funnel so dashboards and Kanban look real.
        for index, application in enumerate(applications):
            score = float(application.ats_score or 0)
            try:
                if score >= 80 and index % 3 == 0:
                    await pipeline.change_status(
                        application,
                        new_status=ApplicationStatus.SHORTLISTED,
                        actor_id=recruiter.id,
                        reason="Strong ATS match",
                        publish_events=False,
                    )
                elif score >= 65 and index % 3 == 1:
                    await pipeline.change_status(
                        application,
                        new_status=ApplicationStatus.UNDER_REVIEW,
                        actor_id=recruiter.id,
                        publish_events=False,
                    )
                elif score < 40 and index % 4 == 0:
                    await pipeline.change_status(
                        application,
                        new_status=ApplicationStatus.REJECTED,
                        actor_id=recruiter.id,
                        reason="Does not meet the core requirements for this role",
                        publish_events=False,
                    )
            except Exception as exc:
                logger.debug("seed_status_skipped", error=str(exc))

        await session.commit()

    # ----------------------------------------------------------- interviews
    async with session_scope() as session:
        shortlisted = list(
            (
                await session.execute(
                    select(Application)
                    .where(
                        Application.company_id == company_id,
                        Application.status == ApplicationStatus.SHORTLISTED,
                    )
                    .options(selectinload(Application.job))
                    .order_by(Application.ats_score.desc())
                    .limit(6)
                )
            )
            .unique()
            .scalars()
            .all()
        )
        interviewer_ids = list(
            (
                await session.execute(
                    select(User.id).where(
                        User.company_id == company_id,
                        User.email.in_(
                            [
                                "rohan.gupta@northwind.test",
                                "anita.shah@northwind.test",
                                "karthik.iyer@northwind.test",
                            ]
                        ),
                    )
                )
            )
            .scalars()
            .all()
        )
        recruiter = await session.scalar(
            select(User).where(User.email == "priya.nair@northwind.test")
        )

        service = InterviewService(session, company_id)
        scheduled: list[uuid.UUID] = []

        for index, application in enumerate(shortlisted):
            # A mix of past interviews (with feedback) and upcoming ones.
            if index < 3:
                start = datetime.now(UTC) - timedelta(days=RANDOM.randint(2, 8), hours=index)
            else:
                start = datetime.now(UTC) + timedelta(days=index - 2, hours=10 + index)
            start = start.replace(minute=0, second=0, microsecond=0)

            try:
                interview = await service.schedule(
                    application_id=application.id,
                    interview_type=(
                        InterviewType.TECHNICAL if index % 2 == 0 else InterviewType.HR
                    ),
                    scheduled_start=start if start > datetime.now(UTC) else
                    datetime.now(UTC) + timedelta(days=1),
                    duration_minutes=60,
                    interviewer_ids=[interviewer_ids[index % len(interviewer_ids)]],
                    organiser_id=recruiter.id,
                    round_name="Technical Round 1" if index % 2 == 0 else "HR Discussion",
                    meeting_link=f"https://meet.northwind.test/{uuid.uuid4().hex[:10]}",
                    send_invitation=True,
                )
                # Backdate the ones that should already have happened.
                if index < 3:
                    interview.scheduled_start = start
                    interview.scheduled_end = start + timedelta(minutes=60)
                scheduled.append(interview.id)
            except Exception as exc:
                logger.debug("seed_interview_skipped", error=str(exc))

        await session.commit()

    # ------------------------------------------------------------- feedback
    async with session_scope() as session:
        from app.models.application import Application
        from app.models.interview import Interview

        past = list(
            (
                await session.execute(
                    select(Interview)
                    .where(
                        Interview.company_id == company_id,
                        Interview.scheduled_start < datetime.now(UTC),
                    )
                    .options(
                        selectinload(Interview.participants),
                        selectinload(Interview.application).selectinload(
                            Application.candidate
                        ),
                        selectinload(Interview.application).selectinload(Application.job),
                    )
                )
            )
            .unique()
            .scalars()
            .all()
        )

        service = InterviewService(session, company_id)
        for index, interview in enumerate(past):
            participant = next((p for p in interview.participants if p.user_id), None)
            if participant is None:
                continue
            recommendation = [
                InterviewRecommendation.STRONG_HIRE,
                InterviewRecommendation.HIRE,
                InterviewRecommendation.MAYBE,
            ][index % 3]
            try:
                await service.submit_feedback(
                    interview,
                    interviewer_id=participant.user_id,
                    overall_rating=[4.5, 4.0, 3.0][index % 3],
                    recommendation=recommendation,
                    technical_skills=[5, 4, 3][index % 3],
                    communication=[4, 4, 4][index % 3],
                    problem_solving=[5, 4, 3][index % 3],
                    domain_knowledge=[4, 4, 3][index % 3],
                    culture_fit=[5, 4, 4][index % 3],
                    strengths=(
                        "- Strong practical knowledge of the core stack\n"
                        "- Communicated design trade-offs clearly\n"
                        "- Good grasp of testing and code review practice"
                    ),
                    weaknesses=(
                        "- Limited hands-on cloud infrastructure experience\n"
                        "- Would benefit from more exposure to large-scale system design"
                    ),
                    comments=(
                        "Solid candidate. Handled the practical exercise well and asked "
                        "good clarifying questions throughout."
                    ),
                    private_remarks="Salary expectation is at the top of our band - worth checking early.",
                )
            except Exception as exc:
                logger.debug("seed_feedback_skipped", error=str(exc))

        await session.commit()

    # --------------------------------------------------------------- offers
    async with session_scope() as session:
        from app.models.application import Application

        top = list(
            (
                await session.execute(
                    select(Application)
                    .where(
                        Application.company_id == company_id,
                        Application.status.in_(
                            [ApplicationStatus.INTERVIEW, ApplicationStatus.SHORTLISTED]
                        ),
                    )
                    .options(
                        selectinload(Application.candidate), selectinload(Application.job)
                    )
                    .order_by(Application.ats_score.desc())
                    .limit(2)
                )
            )
            .unique()
            .scalars()
            .all()
        )
        recruiter = await session.scalar(
            select(User).where(User.email == "priya.nair@northwind.test")
        )
        service = OfferService(session, company_id)

        for application in top:
            try:
                offer = await service.create(
                    application_id=application.id,
                    position_title=application.job.title,
                    base_salary=float(application.job.salary_max or 2500000) * 0.92,
                    created_by_id=recruiter.id,
                    joining_date=date.today() + timedelta(days=45),
                    variable_pay=float(application.job.salary_max or 2500000) * 0.1,
                    joining_bonus=150000,
                    department=None,
                    location=application.job.location_text,
                    employment_type=application.job.employment_type.value,
                    probation_months=3,
                    reporting_to="Vikram Desai",
                    benefits=[
                        "Health insurance for you and your family",
                        "Annual learning budget of INR 75,000",
                        "Employee stock options",
                        "Flexible working hours",
                    ],
                )
                await service.approve(offer, approver_id=recruiter.id)
                await service.send(offer, actor_id=recruiter.id)
            except Exception as exc:
                logger.debug("seed_offer_skipped", error=str(exc))

        await session.commit()


def _print_credentials() -> None:
    line = "=" * 74
    print(f"\n{line}")
    print("  HireHQ demo data is ready")
    print(line)
    print(f"\n  Frontend : {settings.FRONTEND_BASE_URL}")
    print(f"  API docs : {settings.BACKEND_BASE_URL}/docs\n")
    print("  Sign in with any of these accounts:\n")
    print(f"  {'Role':<18} {'Email':<38} Password")
    print(f"  {'-' * 18} {'-' * 38} {'-' * 16}")
    print(
        f"  {'Super Admin':<18} {settings.SEED_SUPER_ADMIN_EMAIL:<38} "
        f"{settings.SEED_SUPER_ADMIN_PASSWORD}"
    )
    shown = {
        "Company Admin": "priya.nair@northwind.test",
        "Recruiter": "arjun.mehta@northwind.test",
        "Hiring Manager": "vikram.desai@northwind.test",
        "Interviewer": "rohan.gupta@northwind.test",
    }
    for label, email in shown.items():
        print(f"  {label:<18} {email:<38} {settings.SEED_DEMO_PASSWORD}")
    print(
        f"  {'Candidate':<18} {'rahul.sharma@example.test':<38} {settings.SEED_DEMO_PASSWORD}"
    )
    print(f"\n{line}\n")


async def _main() -> int:
    parser = argparse.ArgumentParser(description="Seed HireHQ with demo data")
    parser.add_argument(
        "--reset",
        action="store_true",
        help="Delete existing demo data before seeding",
    )
    args = parser.parse_args()

    result = await seed(reset=args.reset)
    if result.get("status") == "already_seeded":
        print(
            "\nDemo data is already present. Re-run with --reset to rebuild it.\n"
        )
        _print_credentials()
        return 0

    _print_credentials()
    print(
        f"  Seeded: {result.get('jobs', 0)} jobs, "
        f"{result.get('candidates', 0)} candidates, "
        f"{result.get('applications', 0)} applications "
        f"({result.get('scored', 0)} ATS-scored)\n"
    )
    return 0


if __name__ == "__main__":
    sys.exit(asyncio.run(_main()))
